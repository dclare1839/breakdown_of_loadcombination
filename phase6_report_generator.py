"""
Phase 6: 자동 리포트 생성
- Word (.docx) : 구조 계산서 형식 보고서
- Excel (.xlsx): 조합별 결과 상세 테이블
- python-docx, openpyxl 사용
"""

from __future__ import annotations
import datetime
from pathlib import Path
from typing import List, Optional
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from phase3_en1990_engine import LoadCombination, EN1990Engine, NA_DATABASE
from phase5_results_visualization import BreakdownAnalyzer, ResultVisualizer


# ─────────────────────────────────────────────
# 헬퍼 : 색상 상수
# ─────────────────────────────────────────────
_BLUE_DARK  = RGBColor(0x0C, 0x44, 0x7C)   # Header
_BLUE_LIGHT = RGBColor(0xB5, 0xD4, 0xF4)   # Sub-header fill
_GRAY_LIGHT = RGBColor(0xF1, 0xEF, 0xE8)   # Row fill
_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

def _hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


# ─────────────────────────────────────────────
# 1. Word 리포트 생성기
# ─────────────────────────────────────────────
class WordReportGenerator:
    """
    EN 1990 Load Combination 검토 보고서 (Word)
    """

    def __init__(self,
                 combos   : List[LoadCombination],
                 analyzer : BreakdownAnalyzer,
                 project_info: Optional[dict] = None):
        self.combos  = combos
        self.ana     = analyzer
        self.doc     = Document()
        self.info    = project_info or {}
        self._setup_styles()

    # ── 스타일 설정 ───────────────────────────
    def _setup_styles(self):
        """기본 스타일 설정"""
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        # 여백 설정 (A4)
        for section in self.doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.0)

    # ── 헬퍼: 제목 단락 ──────────────────────
    def _heading(self, text: str, level: int = 1):
        h = self.doc.add_heading(text, level=level)
        h.runs[0].font.color.rgb = _BLUE_DARK

    # ── 헬퍼: 스타일 테이블 ──────────────────
    def _styled_table(self, headers: List[str], rows: List[List]) -> None:
        """헤더 + 데이터 행 스타일 테이블"""
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 헤더 행
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            run = hdr[i].paragraphs[0].runs[0]
            run.font.bold  = True
            run.font.color.rgb = _WHITE
            run.font.size  = Pt(9)
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc_pr = hdr[i]._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), _hex(_BLUE_DARK))
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            tc_pr.append(shd)

        # 데이터 행
        for ri, row_data in enumerate(rows):
            row = t.add_row().cells
            fill_color = _GRAY_LIGHT if ri % 2 == 0 else _WHITE
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            for ci, val in enumerate(row_data):
                row[ci].text = str(val)
                row[ci].paragraphs[0].runs[0].font.size = Pt(8.5)
                tc_pr = row[ci]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), _hex(fill_color))
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:val"), "clear")
                tc_pr.append(shd)

        self.doc.add_paragraph()  # 테이블 후 공백

    # ── 표지 ─────────────────────────────────
    def _add_cover(self):
        self.doc.add_paragraph()
        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("EN 1990 Load Combination Review")
        run.font.size  = Pt(20)
        run.font.bold  = True
        run.font.color.rgb = _BLUE_DARK

        sub_p = self.doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run("Structural Analysis Report — MIDAS Civil NX")
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(0x5F, 0x5E, 0x5A)

        self.doc.add_paragraph()

        # 프로젝트 정보 테이블
        info_data = [
            ["Project Name", self.info.get("project_name", "—")],
            ["Structure"   , self.info.get("structure",    "—")],
            ["Standard"    , "EN 1990: Basis of Structural Design"],
            ["Country / NA", self.info.get("country_na", "EN (Eurocode default)")],
            ["Prepared by" , self.info.get("prepared_by", "MIDAS Civil NX Plugin")],
            ["Date"        , datetime.date.today().strftime("%d %B %Y")],
        ]
        t = self.doc.add_table(rows=len(info_data), cols=2)
        t.style = "Table Grid"
        for row, (k, v) in zip(t.rows, info_data):
            row.cells[0].text = k
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].text = v
            row.cells[0].width = Cm(4.5)
            row.cells[1].width = Cm(10)

        self.doc.add_page_break()

    # ── 1장: 개요 ─────────────────────────────
    def _add_overview(self):
        self._heading("1. Overview", level=1)
        p = self.doc.add_paragraph(
            "This report presents the automatically generated load combinations "
            "in accordance with EN 1990:2002 (Eurocode – Basis of Structural Design). "
            "The combinations were created using the MIDAS Civil NX EN 1990 Plugin "
            "and cover both Ultimate Limit State (ULS) and Serviceability Limit State (SLS) "
            "verifications."
        )
        p.runs[0].font.size = Pt(10)

        # 조합 통계
        uls_cnt = sum(1 for c in self.combos if c.limit_state == "ULS")
        sls_cnt = sum(1 for c in self.combos if c.limit_state == "SLS")

        self._heading("1.1 Combination Statistics", level=2)
        self._styled_table(
            ["Category", "Count"],
            [["ULS Combinations", uls_cnt],
             ["SLS Combinations", sls_cnt],
             ["Total", len(self.combos)]]
        )

    # ── 2장: EN 1990 조합식 ───────────────────
    def _add_equations(self):
        self._heading("2. Applied Combination Formulae (EN 1990)", level=1)

        eqs = [
            ["ULS STR/GEO", "Eq. 6.10",
             "Σ γG·Gk + γQ·Qk1 + Σ γQ·ψ0i·Qki"],
            ["ULS STR/GEO", "Eq. 6.10a",
             "Σ γG·Gk + Σ γQ·ψ0i·Qki"],
            ["ULS STR/GEO", "Eq. 6.10b",
             "ξ·Σ γG·Gk + γQ·Qk1 + Σ γQ·ψ0i·Qki"],
            ["SLS", "Characteristic",
             "Gk + Qk1 + Σ ψ0i·Qki"],
            ["SLS", "Frequent",
             "Gk + ψ1·Qk1 + Σ ψ2i·Qki"],
            ["SLS", "Quasi-permanent",
             "Gk + Σ ψ2i·Qki"],
        ]
        self._styled_table(["Limit State", "Equation", "Expression"], eqs)

    # ── 3장: 하중 조합 목록 ───────────────────
    def _add_combination_list(self):
        self._heading("3. Generated Load Combinations", level=1)

        rows = []
        for c in self.combos:
            factor_str = "  ".join(
                f"{lc}×{f:.2f}" for lc, f in c.factors.items() if abs(f) > 1e-9
            )
            rows.append([c.name, c.combo_type, c.limit_state, factor_str])

        self._styled_table(
            ["Combination Name", "Type", "LS", "Factors"],
            rows
        )

    # ── 4장: 주요 결과 ────────────────────────
    def _add_key_results(self):
        self._heading("4. Key Results Summary", level=1)

        # 임계 부재
        self._heading("4.1 Critical Elements (Top 10 by Bending Moment)", level=2)
        crit = self.ana.critical_elements(10)
        if not crit.empty:
            rows = [[str(r.ElemID), r.Combo, r.LimitState,
                     f"{r.Axial:.2f}", f"{r.ShearZ:.2f}", f"{r.MomentZ:.2f}"]
                    for r in crit.itertuples()]
            self._styled_table(
                ["ElemID", "Combination", "LS", "N [kN]", "V [kN]", "M [kNm]"],
                rows
            )

        # 최대 변위
        self._heading("4.2 Maximum Displacements (SLS)", level=2)
        crit_d = self.ana.critical_displacements(10)
        if not crit_d.empty:
            rows_d = [[str(r.NodeID), r.Combo,
                       f"{r.Dx:.5f}", f"{r.Dy:.5f}", f"{r.Dz:.5f}", f"{r.D_total:.5f}"]
                      for r in crit_d.itertuples()]
            self._styled_table(
                ["NodeID", "Combination", "Dx [m]", "Dy [m]", "Dz [m]", "D_total [m]"],
                rows_d
            )

    # ── 5장: 그래프 삽입 ─────────────────────
    def _add_charts(self, chart_paths: List[str]):
        if not chart_paths:
            return
        self._heading("5. Result Charts", level=1)
        for path in chart_paths:
            if Path(path).exists():
                self.doc.add_picture(path, width=Cm(15))
                self.doc.add_paragraph()

    # ── 6장: 면책 ─────────────────────────────
    def _add_disclaimer(self):
        self._heading("6. Notes & Disclaimer", level=1)
        self.doc.add_paragraph(
            "This report was automatically generated by the MIDAS Civil NX EN 1990 Plugin. "
            "The engineer of record is responsible for verifying all load combinations, "
            "partial factors, and results against project-specific requirements and the "
            "applicable National Annex. This output does not substitute for professional "
            "engineering judgment."
        )

    # ── 전체 생성 및 저장 ─────────────────────
    def generate(self, output_path: str,
                 chart_paths: Optional[List[str]] = None) -> str:
        self._add_cover()
        self._add_overview()
        self._add_equations()
        self._add_combination_list()
        self._add_key_results()
        self._add_charts(chart_paths or [])
        self._add_disclaimer()

        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        print(f"✅ Word 리포트 저장: {path}")
        return str(path)


# ─────────────────────────────────────────────
# 2. Excel 상세 리포트 생성기
# ─────────────────────────────────────────────
class ExcelReportGenerator:
    """조합별 상세 결과 Excel 리포트"""

    HDR_FILL  = PatternFill("solid", fgColor="0C447C")
    SUB_FILL  = PatternFill("solid", fgColor="B5D4F4")
    ALT_FILL  = PatternFill("solid", fgColor="F1EFE8")
    HDR_FONT  = Font(bold=True, color="FFFFFF", size=10)
    SUB_FONT  = Font(bold=True, color="0C447C", size=9)
    BODY_FONT = Font(size=9)
    CENTER    = Alignment(horizontal="center", vertical="center", wrap_text=True)
    THIN_BD   = Border(
        left  =Side(style="thin"),
        right =Side(style="thin"),
        top   =Side(style="thin"),
        bottom=Side(style="thin")
    )

    def __init__(self, combos: List[LoadCombination], analyzer: BreakdownAnalyzer):
        self.combos  = combos
        self.ana     = analyzer
        self.wb      = Workbook()
        self.wb.remove(self.wb.active)  # 기본 시트 제거

    def _write_sheet(self, ws, df: pd.DataFrame, title: str):
        """DataFrame을 스타일 적용하여 시트에 작성"""
        ws.title = title

        # 헤더
        for ci, col in enumerate(df.columns, 1):
            cell = ws.cell(row=1, column=ci, value=col)
            cell.fill   = self.HDR_FILL
            cell.font   = self.HDR_FONT
            cell.alignment = self.CENTRAL = Alignment(horizontal="center", vertical="center")
            cell.border = self.THIN_BD

        # 데이터
        for ri, row in enumerate(df.itertuples(index=False), 2):
            fill = self.ALT_FILL if ri % 2 == 0 else PatternFill()
            for ci, val in enumerate(row, 1):
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font   = self.BODY_FONT
                cell.fill   = fill
                cell.border = self.THIN_BD
                if isinstance(val, float):
                    cell.number_format = "0.0000"
                    cell.alignment = Alignment(horizontal="right")

        # 열 너비 자동 조정
        for col_cells in ws.columns:
            max_len = max(
                (len(str(c.value)) if c.value else 0) for c in col_cells
            )
            ws.column_dimensions[get_column_letter(col_cells[0].column)].width = \
                min(max_len + 2, 30)

    def generate(self, output_path: str) -> str:
        # 시트 1: Combination List
        ws1 = self.wb.create_sheet("Combinations")
        combo_df = pd.DataFrame([{
            "Name"        : c.name,
            "Type"        : c.combo_type,
            "LimitState"  : c.limit_state,
            "Description" : c.description,
            **{f"F_{k}": round(v, 4) for k, v in c.factors.items()}
        } for c in self.combos])
        self._write_sheet(ws1, combo_df, "Combinations")

        # 시트 2: ULS Envelope
        env_uls = self.ana.force_envelope("ULS")
        if not env_uls.empty:
            ws2 = self.wb.create_sheet("ULS_Envelope")
            self._write_sheet(ws2, env_uls, "ULS_Envelope")

        # 시트 3: SLS Envelope
        env_sls = self.ana.force_envelope("SLS")
        if not env_sls.empty:
            ws3 = self.wb.create_sheet("SLS_Envelope")
            self._write_sheet(ws3, env_sls, "SLS_Envelope")

        # 시트 4: 전체 부재력
        if not self.ana.forces.empty:
            ws4 = self.wb.create_sheet("All_Forces")
            self._write_sheet(ws4, self.ana.forces, "All_Forces")

        # 시트 5: 전체 변위
        if not self.ana.disps.empty:
            ws5 = self.wb.create_sheet("Displacements")
            self._write_sheet(ws5, self.ana.disps, "Displacements")

        # 시트 6: 반력
        if not self.ana.reacts.empty:
            ws6 = self.wb.create_sheet("Reactions")
            self._write_sheet(ws6, self.ana.reacts, "Reactions")

        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.wb.save(str(path))
        print(f"✅ Excel 리포트 저장: {path}")
        return str(path)


# ─────────────────────────────────────────────
# 3. 통합 리포트 생성 함수
# ─────────────────────────────────────────────
def generate_reports(combos    : List[LoadCombination],
                     analyzer  : BreakdownAnalyzer,
                     output_dir: str = "./reports",
                     project_info: Optional[dict] = None) -> dict:
    """
    Word + Excel 리포트 일괄 생성
    반환: {"word": path, "excel": path, "charts": [paths]}
    """
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    # 차트 생성
    viz   = ResultVisualizer(analyzer, output_dir=str(out / "charts"))
    f_png = viz.plot_force_breakdown()
    d_png = viz.plot_displacement_envelope()
    charts = [p for p in [f_png, d_png] if p]

    # Word
    word_gen  = WordReportGenerator(combos, analyzer, project_info)
    word_path = word_gen.generate(
        str(out / "EN1990_LoadCombination_Report.docx"),
        chart_paths=charts
    )

    # Excel
    xls_gen  = ExcelReportGenerator(combos, analyzer)
    xls_path = xls_gen.generate(str(out / "EN1990_Results_Detail.xlsx"))

    print(f"\n📁 리포트 저장 위치: {out.resolve()}")
    return {"word": word_path, "excel": xls_path, "charts": charts}


# ─────────────────────────────────────────────
# 4. 실행
# ─────────────────────────────────────────────
if __name__ == "__main__":
    # ── 샘플 데이터로 리포트 생성 테스트 ──────
    import numpy as np
    from phase3_en1990_engine import EN1990Engine

    sample_lc = pd.DataFrame([
        {"ID": 1, "Name": "SW",   "Type": "ST", "EN1990_Action": "G"},
        {"ID": 2, "Name": "SDL",  "Type": "ST", "EN1990_Action": "G"},
        {"ID": 3, "Name": "LL",   "Type": "LV", "EN1990_Action": "Qi"},
        {"ID": 4, "Name": "WL_X", "Type": "WL", "EN1990_Action": "W"},
        {"ID": 5, "Name": "TMP",  "Type": "TL", "EN1990_Action": "T"},
    ])

    engine = EN1990Engine(sample_lc, country="EN", eq_method="6.10ab")
    combos = engine.generate_all()

    # 가상 결과 데이터 생성 (실제 사용 시 Phase 5 결과 사용)
    rng     = np.random.default_rng(42)
    n_elem  = 20
    n_node  = 15
    elem_ids = list(range(1, n_elem + 1))
    node_ids = list(range(1, n_node + 1))

    force_rows = []
    disp_rows  = []
    react_rows = []

    for combo in combos:
        for eid in elem_ids:
            force_rows.append({
                "ElemID": eid, "Combo": combo.name,
                "ComboType": combo.combo_type, "LimitState": combo.limit_state,
                "Axial"  : float(rng.uniform(-500, 500)),
                "ShearY" : float(rng.uniform(-200, 200)),
                "ShearZ" : float(rng.uniform(-200, 200)),
                "Torque" : float(rng.uniform(-50,  50)),
                "MomentY": float(rng.uniform(-800, 800)),
                "MomentZ": float(rng.uniform(-800, 800)),
            })
        for nid in node_ids:
            dx, dy, dz = rng.uniform(-0.01, 0.01, 3)
            disp_rows.append({
                "NodeID": nid, "Combo": combo.name,
                "ComboType": combo.combo_type, "LimitState": combo.limit_state,
                "Dx": float(dx), "Dy": float(dy), "Dz": float(dz),
                "Rx": 0.0, "Ry": 0.0, "Rz": 0.0,
                "D_total": float(np.sqrt(dx**2 + dy**2 + dz**2))
            })
            react_rows.append({
                "NodeID": nid, "Combo": combo.name,
                "ComboType": combo.combo_type, "LimitState": combo.limit_state,
                "Rx": float(rng.uniform(-100, 100)),
                "Ry": float(rng.uniform(-100, 100)),
                "Rz": float(rng.uniform(0, 500)),
                "Mrx": 0.0, "Mry": 0.0, "Mrz": 0.0,
            })

    results  = {
        "forces"       : pd.DataFrame(force_rows),
        "displacements": pd.DataFrame(disp_rows),
        "reactions"    : pd.DataFrame(react_rows),
    }
    analyzer = BreakdownAnalyzer(results)

    project_info = {
        "project_name": "Sample Bridge - EN 1990 Verification",
        "structure"   : "3-span Composite Bridge",
        "country_na"  : "EN (Eurocode default)",
        "prepared_by" : "MIDAS Civil NX EN1990 Plugin",
    }

    output = generate_reports(
        combos       = combos,
        analyzer     = analyzer,
        output_dir   = "./reports",
        project_info = project_info
    )

    print("\n" + "="*50)
    print("📄 Word :", output["word"])
    print("📊 Excel:", output["excel"])
